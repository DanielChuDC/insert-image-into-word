
# For reading file name
import glob
print(glob.glob('.'))
# For getting the full path
import pathlib
print(pathlib.Path(__file__).parent.absolute())


myList=[]
# initializing split word 
spl_word = '/'
myListWithoutSlash=[]
for file_name in glob.iglob(str(pathlib.Path(__file__).parent.absolute())+ '/imgs/*.jpeg', recursive=True):
  print(file_name)
  myList.append(file_name)
  # using split() 
  # Get String after substring occurrence 
  res = file_name.split('/')
  res= [item for item in res if "jpeg" in item or "JPEG" in item or "png" in item or "PNG" in item or "jpg" in item or "JPG" in item]
  myListWithoutSlash.append(res)

# for python-docs package
# create words documents
from docx import Document
from docx.shared import Cm


document = Document()


table = document.add_table(rows = 1, cols = 4)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ProductName'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'ImageName'
hdr_cells[3].text = 'Image'

for i in range(len(myList)): 
  print(i)
  print(myListWithoutSlash[i])
  row_cells = table.add_row().cells
  row_cells[2].text = myListWithoutSlash[i]
  p = row_cells[3].add_paragraph()
  r = p.add_run()
  r.add_picture(myList[i],width=Cm(4.0), height=Cm(4))


document.save('./output/demo.docx')